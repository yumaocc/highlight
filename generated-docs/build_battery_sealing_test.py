from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = "/Users/q/Desktop/work/highlight/generated-docs/电池封口培训试题_中匈双语.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "B7C3D0"
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, size=11, bold=None, italic=None, color=BLACK, east_asia="Arial Unicode MS", latin="Arial Unicode MS"):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_para_format(paragraph, before=0, after=6, line=1.25, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_mixed_para(parent, parts, before=0, after=6, line=1.25, align=None, style=None):
    paragraph = parent.add_paragraph(style=style)
    set_para_format(paragraph, before=before, after=after, line=line, align=align)
    for text, opts in parts:
        run = paragraph.add_run(text)
        set_run_font(
            run,
            size=opts.get("size", 11),
            bold=opts.get("bold"),
            italic=opts.get("italic"),
            color=opts.get("color", BLACK),
        )
    return paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(int(width * 1440) for width in widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_table(table, widths, header_rows=0):
    set_table_width(table, widths)
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx < header_rows:
                set_cell_shading(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                set_para_format(paragraph, after=2, line=1.15)
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(row_idx < header_rows))


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()


def add_cell_line(cell, text, size=10.5, bold=False, color=BLACK, after=1):
    paragraph = cell.add_paragraph()
    set_para_format(paragraph, after=after, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc, zh, hu, level=1):
    size = 16 if level == 1 else 13
    color = BLUE if level in (1, 2) else DARK_BLUE
    before = 18 if level == 1 else 14
    after = 8 if level == 1 else 6
    paragraph = doc.add_paragraph()
    set_para_format(paragraph, before=before, after=after, line=1.2)
    run = paragraph.add_run(f"{zh} / {hu}")
    set_run_font(run, size=size, bold=True, color=color)
    return paragraph


def add_question(doc, number, zh, hu, options, answer_space=False):
    table = doc.add_table(rows=2, cols=2)
    style_table(table, [0.55, 5.95], header_rows=0)
    table.rows[0].cells[0].merge(table.rows[1].cells[0])
    num_cell = table.rows[0].cells[0]
    clear_cell(num_cell)
    p_num = num_cell.add_paragraph()
    set_para_format(p_num, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_num = p_num.add_run(f"{number}")
    set_run_font(r_num, size=12, bold=True, color=DARK_BLUE)

    q_cell = table.rows[0].cells[1]
    clear_cell(q_cell)
    add_cell_line(q_cell, zh, size=10.5, bold=True, after=1)
    add_cell_line(q_cell, hu, size=10.5, color=GRAY, after=1)

    opt_cell = table.rows[1].cells[1]
    clear_cell(opt_cell)
    if options:
        for zh_opt, hu_opt in options:
            add_cell_line(opt_cell, zh_opt, size=10.2, after=0)
            add_cell_line(opt_cell, hu_opt, size=10.2, color=GRAY, after=2)
    if answer_space:
        add_cell_line(opt_cell, "答 / Válasz: ________________________________", size=10.5, bold=True, after=0)

    spacer = doc.add_paragraph()
    set_para_format(spacer, after=2, line=1)


def add_fill_question(doc, number, zh, hu):
    table = doc.add_table(rows=1, cols=2)
    style_table(table, [0.55, 5.95], header_rows=0)
    num_cell, q_cell = table.rows[0].cells
    clear_cell(num_cell)
    p_num = num_cell.add_paragraph()
    set_para_format(p_num, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_num = p_num.add_run(str(number))
    set_run_font(r_num, size=12, bold=True, color=DARK_BLUE)
    clear_cell(q_cell)
    add_cell_line(q_cell, zh, size=10.5, bold=True, after=2)
    add_cell_line(q_cell, hu, size=10.5, color=GRAY, after=4)
    add_cell_line(q_cell, "答 / Válasz: ________________________________________________________________", size=10.5, after=0)
    spacer = doc.add_paragraph()
    set_para_format(spacer, after=2, line=1)


def add_true_false(doc, number, zh, hu):
    table = doc.add_table(rows=1, cols=3)
    style_table(table, [0.55, 5.35, 0.60], header_rows=0)
    num_cell, q_cell, ans_cell = table.rows[0].cells
    clear_cell(num_cell)
    p_num = num_cell.add_paragraph()
    set_para_format(p_num, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_num = p_num.add_run(str(number))
    set_run_font(r_num, size=12, bold=True, color=DARK_BLUE)
    clear_cell(q_cell)
    add_cell_line(q_cell, zh, size=10.5, bold=True, after=1)
    add_cell_line(q_cell, hu, size=10.5, color=GRAY, after=0)
    clear_cell(ans_cell)
    p_ans = ans_cell.add_paragraph()
    set_para_format(p_ans, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_ans = p_ans.add_run("(    )")
    set_run_font(r_ans, size=11, bold=True)
    spacer = doc.add_paragraph()
    set_para_format(spacer, after=2, line=1)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ):
        style = styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    set_para_format(header, after=0, line=1.0)
    run = header.add_run("电池封口培训试题 / Akkumulátor lezárási képzési teszt")
    set_run_font(run, size=9, color=GRAY)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    set_para_format(footer, after=0, line=1.0)
    run = footer.add_run("95分及以上合格 / Megfelelési határ: legalább 95 pont")
    set_run_font(run, size=9, color=GRAY)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build():
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    set_para_format(title, after=4, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = title.add_run("电池封口培训试题")
    set_run_font(run, size=22, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    set_para_format(subtitle, after=14, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = subtitle.add_run("Akkumulátor lezárási képzési teszt")
    set_run_font(run, size=16, bold=True, color=GRAY)

    meta = doc.add_table(rows=1, cols=2)
    style_table(meta, [4.3, 2.2], header_rows=0)
    for cell in meta.rows[0].cells:
        clear_cell(cell)
        set_cell_shading(cell, LIGHT_GRAY)
    add_cell_line(meta.rows[0].cells[0], "姓名 / Név: ________________________________", size=11, bold=True, after=0)
    add_cell_line(meta.rows[0].cells[1], "总分 / Pontszám: __________", size=11, bold=True, after=0)

    add_heading(doc, "试题说明", "Vizsgaleírás", level=1)
    add_mixed_para(
        doc,
        [
            (
                "本试卷适用于电池装配培训后的考核，考试范围包括电池装配及电池封口操作。试卷满分为100分，95分及以上为合格。考核不合格者有一次补训及补考机会；补考仍不合格者将不予录用。",
                {"size": 10.8},
            )
        ],
        after=4,
    )
    add_mixed_para(
        doc,
        [
            (
                "Ez a teszt az akkumulátor-összeszerelési képzés utáni értékelésre szolgál. A vizsga az akkumulátor-összeszerelési és akkumulátor-lezárási műveletekre vonatkozó ismereteket fedi le. A maximális pontszám 100 pont, a megfelelési határ legalább 95 pont. Sikertelen vizsga esetén egy alkalommal pótképzés és pótvizsga biztosított; ismételt sikertelenség esetén a jelentkező nem alkalmazható.",
                {"size": 10.5, "color": GRAY},
            )
        ],
        after=8,
    )

    add_heading(doc, "一、选择题", "I. Feleletválasztós kérdések", level=1)
    choice_questions = [
        (
            1,
            "根据放入物体的（     ），选择使用手套箱的大、小真空仓。",
            "A behelyezendő tárgy (     ) alapján kell kiválasztani a kesztyűs doboz nagy vagy kis vákuumkamráját.",
            [
                ("A. 大小", "A. mérete"),
                ("B. 颜色", "B. színe"),
                ("C. 材质", "C. anyaga"),
                ("D. 表面光滑程度", "D. felületi simasága"),
            ],
        ),
        (
            2,
            "从手套箱中取出或向手套箱中放入物品结束后，真空仓上方压力表读数应为（     ），且保持不变。",
            "A tárgy kesztyűs dobozból történő kivétele vagy behelyezése után a vákuumkamra felső nyomásmérőjének (     ) értéket kell mutatnia, és ennek változatlannak kell maradnia.",
            [("A. 0", "A. 0"), ("B. 1", "B. 1"), ("C. -1", "C. -1"), ("D. 2", "D. 2")],
        ),
        (
            3,
            "从手套箱中取出或向手套箱中放入物品结束后，真空仓下方旋转阀应处于（     ）位置，且保持不变。",
            "A tárgy kesztyűs dobozból történő kivétele vagy behelyezése után a vákuumkamra alsó forgószelepének (     ) állásban kell lennie, és ebben kell maradnia.",
            [
                ("A. 抽真空", "A. vákuumozás"),
                ("B. 补气", "B. gázfeltöltés"),
                ("C. 关闭", "C. zárva"),
                ("D. 任一位置", "D. bármelyik állás"),
            ],
        ),
        (
            4,
            "扣电封口机可通过调节封口（     ）和压力，调节电池封口的松紧程度。调试完成后，这两项参数不得随意改变。",
            "A gombelem-lezáró gépen a lezárási (     ) és a nyomás beállításával szabályozható a lezárás szorossága. A beállítás befejezése után ezt a két paramétert nem szabad önkényesen módosítani.",
            [
                ("A. 时间", "A. idő"),
                ("B. 距离", "B. távolság"),
                ("C. 角度", "C. szög"),
                ("D. 水平", "D. vízszint"),
            ],
        ),
        (
            5,
            "电池封口完成后，需使用无水乙醇对电池表面进行清洁，防止溢出的（     ）腐蚀设备。",
            "A lezárás befejezése után a cella felületét vízmentes etanollal meg kell tisztítani, hogy a kifolyt (     ) ne marja meg a berendezést.",
            [
                ("A. 酒精", "A. alkohol"),
                ("B. NMP", "B. NMP"),
                ("C. 电解液", "C. elektrolit"),
                ("D. 水", "D. víz"),
            ],
        ),
        (
            6,
            "电池封口时应（     ）向上放置。",
            "Lezáráskor az akkumulátort úgy kell elhelyezni, hogy a(z) (     ) felfelé nézzen.",
            [
                ("A. 正极", "A. pozitív pólus"),
                ("B. 负极", "B. negatív pólus"),
                ("C. 立放", "C. állítva"),
                ("D. 都可以", "D. mindegy"),
            ],
        ),
        (
            7,
            "发生紧急情况时，应立即按下（     ）按钮。",
            "Vészhelyzet esetén azonnal meg kell nyomni a(z) (     ) gombot.",
            [
                ("A. 应急", "A. vészleállító"),
                ("B. 上升", "B. fel"),
                ("C. 下降", "C. le"),
                ("D. 暂停", "D. szünet"),
            ],
        ),
        (
            8,
            "封口模具大小为（     ）。",
            "A lezáró szerszám mérete: (     ).",
            [("A. 2430", "A. 2430"), ("B. 2032", "B. 2032"), ("C. 1632", "C. 1632"), ("D. 2450", "D. 2450")],
        ),
    ]
    for q in choice_questions:
        add_question(doc, *q, answer_space=True)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "二、填空题", "II. Kiegészítendő kérdések", level=1)
    fill_questions = [
        (
            1,
            "电池封口过程中，应在电池上垫一层_____，目的是_____。",
            "Lezárás közben egy réteg _____ kell az akkumulátorra helyezni; ennek célja: _____.",
        ),
        (
            2,
            "夹取电池时，应使用_____。",
            "Az akkumulátor megfogásához _____ kell használni.",
        ),
        (
            3,
            "电池模具应在_____、_____使用酒精进行擦拭。",
            "A cellaformát _____ és _____ alkohollal le kell törölni.",
        ),
        (
            4,
            "封口机使用完毕后应_____。",
            "Használat után a lezárógépet _____.",
        ),
    ]
    for q in fill_questions:
        add_fill_question(doc, *q)

    add_heading(doc, "三、判断题", "III. Igaz vagy hamis", level=1)
    tf_questions = [
        (
            1,
            "封口过程中，操作人员应与封口机保持一定距离。",
            "Lezárás közben a kezelőnek megfelelő távolságot kell tartania a lezárógéptől.",
        ),
        (
            2,
            "封口完毕后需用大量酒精擦拭电池。",
            "Lezárás után nagy mennyiségű alkohollal kell letörölni az akkumulátort.",
        ),
        (
            3,
            "电池从过渡仓取出前，不可在真空仓内直接对电池抽真空。",
            "Mielőtt az akkumulátort kivenné az átadókamrából, tilos a cellát közvetlenül vákuum alá helyezni a vákuumkamrában.",
        ),
        (
            4,
            "电池封口过程中可以不垫无纺布。",
            "Lezárás közben nem szükséges nemszőtt kendőt használni.",
        ),
        (
            5,
            "电池不易用镊子夹取时，可以直接用手拿取电池。",
            "Ha az akkumulátort nehéz csipesszel megfogni, kézzel is megfogható.",
        ),
        (
            6,
            "操作人员手部可以随意放置。",
            "A kezelő a kezét munka közben bárhová szabadon elhelyezheti.",
        ),
    ]
    for q in tf_questions:
        add_true_false(doc, *q)

    note = doc.add_table(rows=1, cols=1)
    style_table(note, [6.5], header_rows=0)
    set_cell_shading(note.rows[0].cells[0], LIGHT_GRAY)
    clear_cell(note.rows[0].cells[0])
    add_cell_line(note.rows[0].cells[0], "备注 / Megjegyzés:", size=10.5, bold=True, after=4)
    add_cell_line(note.rows[0].cells[0], "____________________________________________________________________________", size=10.5, after=4)
    add_cell_line(note.rows[0].cells[0], "____________________________________________________________________________", size=10.5, after=0)

    doc.core_properties.title = "电池封口培训试题 / Akkumulátor lezárási képzési teszt"
    doc.core_properties.subject = "Battery sealing training test"
    doc.core_properties.author = "Codex"
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
